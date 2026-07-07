"""
doc_generator.py
-----------------
Takes the agent's structured output (title, assumptions, sections) and
renders a polished Word document using python-docx.
"""

import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = "generated_docs"


def _add_bullets_or_paragraph(doc: Document, text: str):
    """
    If the LLM returned bullet-style lines (starting with -, *, or •),
    render them as a real bullet list. Otherwise render as a normal paragraph.
    """
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    bullet_lines = [l for l in lines if re.match(r"^[-*•]\s+", l)]

    if len(bullet_lines) >= 2:
        for line in lines:
            clean = re.sub(r"^[-*•]\s+", "", line)
            doc.add_paragraph(clean, style="List Bullet")
    else:
        doc.add_paragraph(text)


def generate_docx(agent_output: dict, request_id: str) -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()

    # --- Title page / header ---
    title = doc.add_heading(agent_output["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Document type: {agent_output['document_type'].replace('_', ' ').title()}\n"
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.italic = True
    run.font.size = Pt(10)

    # --- Assumptions (only shown if the agent had to make any) ---
    if agent_output.get("assumptions"):
        doc.add_heading("Assumptions Made by the Agent", level=1)
        for a in agent_output["assumptions"]:
            doc.add_paragraph(a, style="List Bullet")

    # --- Body sections ---
    for section in agent_output["sections"]:
        doc.add_heading(section["section_title"], level=1)
        _add_bullets_or_paragraph(doc, section["content"])

    filename = f"{OUTPUT_DIR}/{request_id}.docx"
    doc.save(filename)
    return filename
